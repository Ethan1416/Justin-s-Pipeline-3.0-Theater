"""
Document Generation Agents (HARDCODED)
======================================

Specialized agents for converting markdown files to Word documents
in the production folder.

All conversion logic is HARDCODED for consistent document formatting.
"""

from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass, field
import re

from .base import Agent, AgentResult, AgentStatus

# Try to import python-docx, handle if not installed
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# =============================================================================
# HARDCODED DOCUMENT STYLES
# =============================================================================

# HARDCODED: Document formatting settings
DOCUMENT_STYLES = {
    "title": {
        "font_size": 24,
        "bold": True,
        "color": (0, 51, 102),  # Dark blue
    },
    "heading1": {
        "font_size": 18,
        "bold": True,
        "color": (0, 51, 102),
    },
    "heading2": {
        "font_size": 14,
        "bold": True,
        "color": (51, 51, 51),
    },
    "heading3": {
        "font_size": 12,
        "bold": True,
        "color": (51, 51, 51),
    },
    "body": {
        "font_size": 11,
        "bold": False,
        "color": (0, 0, 0),
    },
    "bullet": {
        "font_size": 11,
        "bold": False,
        "color": (0, 0, 0),
    },
}

# HARDCODED: File type mappings for naming
FILE_TYPE_TITLES = {
    "lesson_plan": "Lesson Plan",
    "exit_ticket": "Exit Ticket",
    "journal_prompts": "Journal Prompts",
    "handout_activity": "Activity Handout",
    "unit_plan": "Unit Plan",
}


# =============================================================================
# MarkdownParserAgent
# =============================================================================

class MarkdownParserAgent(Agent):
    """
    HARDCODED agent for parsing markdown files into structured content.

    Parses markdown syntax into structured blocks for document generation.
    """

    def __init__(self):
        super().__init__(name="MarkdownParserAgent")

    def _process(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Parse markdown content into structured blocks."""
        md_content = context.get("markdown_content", "")
        file_path = context.get("file_path", "")

        if not md_content:
            return {"blocks": [], "error": "No markdown content provided"}

        blocks = self._parse_markdown(md_content)

        # Extract title from first heading or filename
        title = self._extract_title(blocks, file_path)

        return {
            "title": title,
            "blocks": blocks,
            "block_count": len(blocks),
            "source_file": str(file_path),
        }

    def _parse_markdown(self, content: str) -> List[Dict[str, Any]]:
        """
        HARDCODED markdown parsing logic.

        Supports:
        - # Heading 1, ## Heading 2, ### Heading 3
        - **bold**, *italic*
        - - bullet points
        - 1. numbered lists
        - > blockquotes
        - ``` code blocks
        - --- horizontal rules
        - Regular paragraphs
        """
        blocks = []
        lines = content.split('\n')

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Skip empty lines
            if not stripped:
                i += 1
                continue

            # Code block
            if stripped.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                blocks.append({
                    "type": "code",
                    "content": '\n'.join(code_lines),
                })
                i += 1
                continue

            # Heading 1
            if stripped.startswith('# ') and not stripped.startswith('##'):
                blocks.append({
                    "type": "heading1",
                    "content": stripped[2:].strip(),
                })
                i += 1
                continue

            # Heading 2
            if stripped.startswith('## ') and not stripped.startswith('###'):
                blocks.append({
                    "type": "heading2",
                    "content": stripped[3:].strip(),
                })
                i += 1
                continue

            # Heading 3
            if stripped.startswith('### '):
                blocks.append({
                    "type": "heading3",
                    "content": stripped[4:].strip(),
                })
                i += 1
                continue

            # Horizontal rule
            if stripped in ['---', '***', '___']:
                blocks.append({
                    "type": "horizontal_rule",
                    "content": "",
                })
                i += 1
                continue

            # Blockquote
            if stripped.startswith('>'):
                quote_lines = [stripped[1:].strip()]
                i += 1
                while i < len(lines) and lines[i].strip().startswith('>'):
                    quote_lines.append(lines[i].strip()[1:].strip())
                    i += 1
                blocks.append({
                    "type": "blockquote",
                    "content": ' '.join(quote_lines),
                })
                continue

            # Bullet list
            if stripped.startswith('- ') or stripped.startswith('* '):
                bullet_items = [stripped[2:].strip()]
                i += 1
                while i < len(lines):
                    next_line = lines[i].strip()
                    if next_line.startswith('- ') or next_line.startswith('* '):
                        bullet_items.append(next_line[2:].strip())
                        i += 1
                    elif next_line.startswith('  ') and bullet_items:
                        # Continuation of previous item
                        bullet_items[-1] += ' ' + next_line.strip()
                        i += 1
                    else:
                        break
                blocks.append({
                    "type": "bullet_list",
                    "items": bullet_items,
                })
                continue

            # Numbered list
            if re.match(r'^\d+\.\s', stripped):
                numbered_items = [re.sub(r'^\d+\.\s*', '', stripped)]
                i += 1
                while i < len(lines):
                    next_line = lines[i].strip()
                    if re.match(r'^\d+\.\s', next_line):
                        numbered_items.append(re.sub(r'^\d+\.\s*', '', next_line))
                        i += 1
                    else:
                        break
                blocks.append({
                    "type": "numbered_list",
                    "items": numbered_items,
                })
                continue

            # Regular paragraph
            para_lines = [stripped]
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line or next_line.startswith('#') or next_line.startswith('-') or \
                   next_line.startswith('*') or next_line.startswith('>') or \
                   next_line.startswith('```') or re.match(r'^\d+\.\s', next_line):
                    break
                para_lines.append(next_line)
                i += 1

            blocks.append({
                "type": "paragraph",
                "content": ' '.join(para_lines),
            })

        return blocks

    def _extract_title(self, blocks: List[Dict], file_path: str) -> str:
        """Extract title from first heading or generate from filename."""
        for block in blocks:
            if block["type"] == "heading1":
                return block["content"]

        # Generate from filename
        if file_path:
            filename = Path(file_path).stem
            # Convert snake_case to Title Case
            title = filename.replace('_', ' ').title()
            # Check for known file types
            for key, value in FILE_TYPE_TITLES.items():
                if key in filename.lower():
                    return value
            return title

        return "Document"


# =============================================================================
# MarkdownToWordAgent
# =============================================================================

class MarkdownToWordAgent(Agent):
    """
    HARDCODED agent for converting parsed markdown to Word documents.

    Uses python-docx to generate professionally formatted Word documents.
    """

    def __init__(self):
        super().__init__(name="MarkdownToWordAgent")

    def _process(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Convert parsed markdown blocks to Word document."""
        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "python-docx not installed. Run: pip install python-docx"
            }

        blocks = context.get("blocks", [])
        title = context.get("title", "Document")
        output_path = context.get("output_path", "")

        if not blocks:
            return {"success": False, "error": "No content blocks provided"}

        if not output_path:
            return {"success": False, "error": "No output path provided"}

        # Create document
        doc = Document()

        # Add title
        self._add_title(doc, title)

        # Process each block
        for block in blocks:
            self._add_block(doc, block)

        # Save document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        return {
            "success": True,
            "output_path": str(output_path),
            "title": title,
            "block_count": len(blocks),
        }

    def _add_title(self, doc: Document, title: str):
        """Add formatted title to document."""
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(DOCUMENT_STYLES["title"]["font_size"])
        run.font.color.rgb = RGBColor(*DOCUMENT_STYLES["title"]["color"])
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Add spacing

    def _add_block(self, doc: Document, block: Dict[str, Any]):
        """Add a content block to the document."""
        block_type = block.get("type", "paragraph")

        if block_type == "heading1":
            self._add_heading(doc, block["content"], 1)

        elif block_type == "heading2":
            self._add_heading(doc, block["content"], 2)

        elif block_type == "heading3":
            self._add_heading(doc, block["content"], 3)

        elif block_type == "paragraph":
            self._add_paragraph(doc, block["content"])

        elif block_type == "bullet_list":
            for item in block.get("items", []):
                self._add_bullet(doc, item)

        elif block_type == "numbered_list":
            for i, item in enumerate(block.get("items", []), 1):
                self._add_numbered_item(doc, item, i)

        elif block_type == "blockquote":
            self._add_blockquote(doc, block["content"])

        elif block_type == "code":
            self._add_code_block(doc, block["content"])

        elif block_type == "horizontal_rule":
            self._add_horizontal_rule(doc)

    def _add_heading(self, doc: Document, text: str, level: int):
        """Add a heading to the document."""
        style_key = f"heading{level}"
        style = DOCUMENT_STYLES.get(style_key, DOCUMENT_STYLES["heading1"])

        para = doc.add_paragraph()
        run = para.add_run(self._clean_markdown(text))
        run.bold = style["bold"]
        run.font.size = Pt(style["font_size"])
        run.font.color.rgb = RGBColor(*style["color"])

    def _add_paragraph(self, doc: Document, text: str):
        """Add a paragraph with inline formatting."""
        para = doc.add_paragraph()
        self._add_formatted_text(para, text)

    def _add_bullet(self, doc: Document, text: str):
        """Add a bullet point."""
        para = doc.add_paragraph(style='List Bullet')
        self._add_formatted_text(para, text)

    def _add_numbered_item(self, doc: Document, text: str, number: int):
        """Add a numbered list item."""
        para = doc.add_paragraph(style='List Number')
        self._add_formatted_text(para, text)

    def _add_blockquote(self, doc: Document, text: str):
        """Add a blockquote."""
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        run = para.add_run(self._clean_markdown(text))
        run.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)

    def _add_code_block(self, doc: Document, text: str):
        """Add a code block."""
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        run = para.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    def _add_horizontal_rule(self, doc: Document):
        """Add a horizontal rule (represented as underscores)."""
        para = doc.add_paragraph()
        para.add_run('_' * 50)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_formatted_text(self, para, text: str):
        """Add text with bold/italic formatting from markdown."""
        # Process **bold** and *italic*
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
        parts = re.findall(pattern, text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = para.add_run(part[1:-1])
                run.italic = True
            else:
                para.add_run(part)

    def _clean_markdown(self, text: str) -> str:
        """Remove markdown formatting from text."""
        # Remove bold
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        # Remove italic
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        # Remove links [text](url) -> text
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        return text


# =============================================================================
# ProductionDocGeneratorAgent
# =============================================================================

class ProductionDocGeneratorAgent(Agent):
    """
    HARDCODED orchestrator agent for converting all markdown files
    in the production folder to Word documents.
    """

    # HARDCODED: Production folder path
    PRODUCTION_PATH = Path(__file__).parent.parent / "production"

    def __init__(self):
        super().__init__(name="ProductionDocGeneratorAgent")
        self.parser = MarkdownParserAgent()
        self.converter = MarkdownToWordAgent()

    def _process(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Find all markdown files and convert to Word documents."""
        production_path = context.get("production_path", self.PRODUCTION_PATH)
        production_path = Path(production_path)

        if not production_path.exists():
            return {
                "success": False,
                "error": f"Production path not found: {production_path}"
            }

        # Find all markdown files
        md_files = list(production_path.rglob("*.md"))

        if not md_files:
            return {
                "success": False,
                "error": "No markdown files found in production folder"
            }

        results = []
        successful = 0
        failed = 0

        for md_file in md_files:
            result = self._convert_file(md_file)
            results.append(result)
            if result.get("success"):
                successful += 1
            else:
                failed += 1

        return {
            "success": failed == 0,
            "total_files": len(md_files),
            "successful": successful,
            "failed": failed,
            "conversions": results,
        }

    def _convert_file(self, md_file: Path) -> Dict[str, Any]:
        """Convert a single markdown file to Word document."""
        try:
            # Read markdown content
            with open(md_file, 'r', encoding='utf-8') as f:
                md_content = f.read()

            # Parse markdown
            parse_result = self.parser.execute({
                "markdown_content": md_content,
                "file_path": str(md_file),
            })

            if parse_result.status != AgentStatus.COMPLETED:
                return {
                    "success": False,
                    "source": str(md_file),
                    "error": f"Parse failed: {parse_result.errors}"
                }

            # Generate output path (same location, .docx extension)
            output_path = md_file.with_suffix('.docx')

            # Convert to Word
            convert_result = self.converter.execute({
                "blocks": parse_result.output.get("blocks", []),
                "title": parse_result.output.get("title", "Document"),
                "output_path": str(output_path),
            })

            if convert_result.status != AgentStatus.COMPLETED:
                return {
                    "success": False,
                    "source": str(md_file),
                    "error": f"Conversion failed: {convert_result.errors}"
                }

            return {
                "success": True,
                "source": str(md_file),
                "output": str(output_path),
                "title": parse_result.output.get("title"),
            }

        except Exception as e:
            return {
                "success": False,
                "source": str(md_file),
                "error": str(e)
            }


# =============================================================================
# FileScannerAgent
# =============================================================================

class FileScannerAgent(Agent):
    """
    HARDCODED agent for scanning production folder and identifying
    markdown files that need conversion.
    """

    def __init__(self):
        super().__init__(name="FileScannerAgent")

    def _process(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Scan production folder for markdown files."""
        production_path = context.get("production_path", "")

        if not production_path:
            return {"success": False, "error": "No production path provided"}

        production_path = Path(production_path)

        if not production_path.exists():
            return {"success": False, "error": f"Path not found: {production_path}"}

        # Find all markdown files
        md_files = list(production_path.rglob("*.md"))

        # Check which already have Word equivalents
        files_info = []
        needs_conversion = []
        already_exists = []

        for md_file in md_files:
            docx_path = md_file.with_suffix('.docx')
            info = {
                "markdown_path": str(md_file),
                "word_path": str(docx_path),
                "word_exists": docx_path.exists(),
                "relative_path": str(md_file.relative_to(production_path)),
            }
            files_info.append(info)

            if docx_path.exists():
                already_exists.append(str(md_file))
            else:
                needs_conversion.append(str(md_file))

        return {
            "success": True,
            "total_markdown_files": len(md_files),
            "needs_conversion": len(needs_conversion),
            "already_converted": len(already_exists),
            "files": files_info,
            "files_to_convert": needs_conversion,
        }

"""
Document Parser - Generic document parsing for NCLEX pipeline.

This module provides generic document parsing capabilities for various
file formats used in the pipeline. It serves as a unified interface
for parsing different document types.

Supports:
- .docx (Word documents)
- .txt (Plain text)
- .md (Markdown)
- .pdf (PDF documents - basic support)

Usage:
    from skills.parsing.document_parser import DocumentParser

    parser = DocumentParser()
    result = parser.parse("path/to/document.docx")
"""

import re
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass, field


@dataclass
class Section:
    """Represents a document section."""
    title: str
    content: str
    level: int = 1
    start_line: int = 0
    style: str = "Normal"


@dataclass
class ParsedDocument:
    """Container for parsed document content."""
    file_path: str
    file_name: str
    file_type: str
    success: bool = False
    raw_text: str = ""
    sections: List[Section] = field(default_factory=list)
    paragraphs: List[str] = field(default_factory=list)
    headings: List[Dict] = field(default_factory=list)
    word_count: int = 0
    line_count: int = 0
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    metadata: Dict = field(default_factory=dict)


class DocumentParser:
    """Generic document parser for multiple formats."""

    # Supported file extensions
    SUPPORTED_EXTENSIONS = ['.docx', '.txt', '.md', '.pdf']

    # Heading patterns for Markdown/text
    HEADING_PATTERNS = [
        (r'^#\s+(.+)$', 1),      # H1
        (r'^##\s+(.+)$', 2),     # H2
        (r'^###\s+(.+)$', 3),    # H3
        (r'^####\s+(.+)$', 4),   # H4
        (r'^(.+)\n=+$', 1),      # Underlined H1
        (r'^(.+)\n-+$', 2),      # Underlined H2
    ]

    def __init__(self):
        """Initialize the DocumentParser."""
        self._check_dependencies()

    def _check_dependencies(self):
        """Check for available parsing libraries."""
        self._has_docx = False
        self._has_pdf = False

        try:
            import docx
            self._has_docx = True
        except ImportError:
            pass

        try:
            import PyPDF2
            self._has_pdf = True
        except ImportError:
            pass

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse a document file.

        Args:
            file_path: Path to document

        Returns:
            ParsedDocument with parsed content
        """
        path = Path(file_path)
        result = ParsedDocument(
            file_path=str(path),
            file_name=path.name,
            file_type=path.suffix.lower(),
            success=False
        )

        if not path.exists():
            result.errors.append(f"File not found: {file_path}")
            return result

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            result.errors.append(f"Unsupported file format: {extension}")
            result.errors.append(f"Supported formats: {', '.join(self.SUPPORTED_EXTENSIONS)}")
            return result

        # Route to appropriate parser
        if extension == '.docx':
            return self._parse_docx(path, result)
        elif extension in ['.txt', '.md']:
            return self._parse_text(path, result)
        elif extension == '.pdf':
            return self._parse_pdf(path, result)

        return result

    def parse_text_content(
        self,
        text: str,
        source_name: str = "text_input"
    ) -> ParsedDocument:
        """
        Parse raw text content.

        Args:
            text: Text to parse
            source_name: Name for the source

        Returns:
            ParsedDocument with parsed content
        """
        result = ParsedDocument(
            file_path=source_name,
            file_name=source_name,
            file_type=".txt",
            success=False
        )

        try:
            result.raw_text = text
            result.paragraphs = self._extract_paragraphs(text)
            result.sections = self._extract_sections(text)
            result.headings = self._extract_headings(text)
            result.word_count = len(text.split())
            result.line_count = len(text.split('\n'))
            result.success = True
        except Exception as e:
            result.errors.append(f"Parse error: {e}")

        return result

    def _parse_docx(self, path: Path, result: ParsedDocument) -> ParsedDocument:
        """Parse a Word document."""
        if not self._has_docx:
            result.errors.append("python-docx not installed. Run: pip install python-docx")
            return result

        try:
            from docx import Document
            doc = Document(str(path))

            # Extract paragraphs with styles
            paragraphs = []
            headings = []
            raw_lines = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
                    raw_lines.append(text)

                    style_name = para.style.name if para.style else 'Normal'
                    if 'Heading' in style_name:
                        # Extract heading level from style name
                        level_match = re.search(r'Heading\s*(\d+)', style_name)
                        level = int(level_match.group(1)) if level_match else 1
                        headings.append({
                            'text': text,
                            'level': level,
                            'style': style_name
                        })

            result.raw_text = '\n'.join(raw_lines)
            result.paragraphs = paragraphs
            result.headings = headings
            result.sections = self._build_sections_from_headings(headings, paragraphs)
            result.word_count = len(result.raw_text.split())
            result.line_count = len(paragraphs)

            # Extract metadata
            try:
                props = doc.core_properties
                result.metadata = {
                    'author': props.author or '',
                    'title': props.title or '',
                    'created': str(props.created) if props.created else '',
                    'modified': str(props.modified) if props.modified else ''
                }
            except Exception:
                pass

            result.success = True

        except Exception as e:
            result.errors.append(f"Error parsing DOCX: {e}")

        return result

    def _parse_text(self, path: Path, result: ParsedDocument) -> ParsedDocument:
        """Parse a text or markdown file."""
        try:
            # Try UTF-8 first
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    text = f.read()
            except UnicodeDecodeError:
                with open(path, 'r', encoding='latin-1') as f:
                    text = f.read()
                result.warnings.append("File read with latin-1 encoding")

            result.raw_text = text
            result.paragraphs = self._extract_paragraphs(text)
            result.sections = self._extract_sections(text)
            result.headings = self._extract_headings(text)
            result.word_count = len(text.split())
            result.line_count = len(text.split('\n'))

            # File metadata
            stat = path.stat()
            result.metadata = {
                'size_bytes': stat.st_size,
                'modified': stat.st_mtime
            }

            result.success = True

        except Exception as e:
            result.errors.append(f"Error reading file: {e}")

        return result

    def _parse_pdf(self, path: Path, result: ParsedDocument) -> ParsedDocument:
        """Parse a PDF document (basic text extraction)."""
        if not self._has_pdf:
            result.errors.append("PyPDF2 not installed. Run: pip install PyPDF2")
            return result

        try:
            import PyPDF2

            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)

                text_parts = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                text = '\n\n'.join(text_parts)

            result.raw_text = text
            result.paragraphs = self._extract_paragraphs(text)
            result.sections = self._extract_sections(text)
            result.word_count = len(text.split())
            result.line_count = len(text.split('\n'))

            # PDF metadata
            result.metadata = {
                'pages': len(reader.pages)
            }

            result.success = True
            result.warnings.append("PDF parsing is basic - formatting may be lost")

        except Exception as e:
            result.errors.append(f"Error parsing PDF: {e}")

        return result

    def _extract_paragraphs(self, text: str) -> List[str]:
        """Extract paragraphs from text."""
        # Split by double newlines or other paragraph markers
        raw_paragraphs = re.split(r'\n\s*\n', text)
        return [p.strip() for p in raw_paragraphs if p.strip()]

    def _extract_headings(self, text: str) -> List[Dict]:
        """Extract headings from text using patterns."""
        headings = []
        lines = text.split('\n')

        for i, line in enumerate(lines):
            for pattern, level in self.HEADING_PATTERNS:
                match = re.match(pattern, line)
                if match:
                    headings.append({
                        'text': match.group(1).strip(),
                        'level': level,
                        'line': i
                    })
                    break

        return headings

    def _extract_sections(self, text: str) -> List[Section]:
        """Extract sections from text based on headings."""
        headings = self._extract_headings(text)
        if not headings:
            # No headings - treat as single section
            return [Section(title="Document", content=text, level=1)]

        lines = text.split('\n')
        sections = []

        for i, heading in enumerate(headings):
            start_line = heading['line']
            end_line = headings[i + 1]['line'] if i + 1 < len(headings) else len(lines)

            content_lines = lines[start_line + 1:end_line]
            content = '\n'.join(content_lines).strip()

            sections.append(Section(
                title=heading['text'],
                content=content,
                level=heading['level'],
                start_line=start_line
            ))

        return sections

    def _build_sections_from_headings(
        self,
        headings: List[Dict],
        paragraphs: List[str]
    ) -> List[Section]:
        """Build sections from DOCX headings and paragraphs."""
        if not headings:
            return [Section(title="Document", content='\n'.join(paragraphs))]

        sections = []
        heading_texts = {h['text'] for h in headings}

        current_section = None
        content_buffer = []

        for para in paragraphs:
            if para in heading_texts:
                # Save previous section
                if current_section:
                    current_section.content = '\n'.join(content_buffer)
                    sections.append(current_section)

                # Start new section
                heading_info = next((h for h in headings if h['text'] == para), None)
                current_section = Section(
                    title=para,
                    content='',
                    level=heading_info['level'] if heading_info else 1
                )
                content_buffer = []
            else:
                content_buffer.append(para)

        # Save last section
        if current_section:
            current_section.content = '\n'.join(content_buffer)
            sections.append(current_section)

        return sections

    def format_report(self, result: ParsedDocument) -> str:
        """Format parse result as report."""
        lines = [
            "=" * 60,
            "DOCUMENT PARSE REPORT",
            "=" * 60,
            f"File: {result.file_name}",
            f"Type: {result.file_type}",
            f"Status: {'SUCCESS' if result.success else 'FAILED'}",
            "",
            f"Words: {result.word_count}",
            f"Lines: {result.line_count}",
            f"Paragraphs: {len(result.paragraphs)}",
            f"Sections: {len(result.sections)}",
            f"Headings: {len(result.headings)}",
            ""
        ]

        if result.errors:
            lines.append("ERRORS:")
            for error in result.errors:
                lines.append(f"  - {error}")
            lines.append("")

        if result.warnings:
            lines.append("WARNINGS:")
            for warn in result.warnings:
                lines.append(f"  - {warn}")
            lines.append("")

        if result.headings:
            lines.append("-" * 60)
            lines.append("HEADINGS:")
            for heading in result.headings[:10]:
                indent = "  " * heading.get('level', 1)
                lines.append(f"{indent}{heading['text'][:50]}")

        if result.sections:
            lines.append("")
            lines.append("-" * 60)
            lines.append("SECTIONS:")
            for section in result.sections[:5]:
                lines.append(f"  Level {section.level}: {section.title[:40]}")
                preview = section.content[:80].replace('\n', ' ')
                lines.append(f"    {preview}...")

        lines.append("")
        lines.append("=" * 60)
        return "\n".join(lines)


def parse_document(file_path: str) -> ParsedDocument:
    """Convenience function to parse a document."""
    parser = DocumentParser()
    return parser.parse(file_path)


if __name__ == "__main__":
    import sys

    print("Document Parser - NCLEX Pipeline Parsing Skill")
    print("=" * 50)

    parser = DocumentParser()

    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        print(f"\nParsing: {file_path}")
        result = parser.parse(file_path)
    else:
        # Demo with sample text
        sample_text = """
# Introduction

This is the introduction section with some content
about the document's purpose.

## Section One

Content for section one goes here.
It can span multiple paragraphs.

More content in section one.

## Section Two

Content for section two.

### Subsection 2.1

Detailed content in a subsection.

## Conclusion

Final thoughts and summary.
"""

        print("\nDemo mode - parsing sample markdown text")
        result = parser.parse_text_content(sample_text, "demo_document.md")

    report = parser.format_report(result)
    print(report)

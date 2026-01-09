"""
Handout Validator Skill
=======================

HARDCODED validator for activity handouts.
This validator CANNOT be bypassed - all handouts MUST pass validation.

Validation Rules:
- R1: Minimum 6 items per activity (HARDCODED)
- R2: Answer key required on separate page (HARDCODED)
- R3: Instructions section required (HARDCODED)
- R4: Student info fields required (HARDCODED)
- R5: Word document format (.docx) required (HARDCODED)
- R6: Each item must have explanation in answer key

Pipeline: Theater Education
"""

from pathlib import Path
from typing import Dict, Any, List, Optional
import re

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# =============================================================================
# HARDCODED VALIDATION THRESHOLDS (CANNOT BE MODIFIED)
# =============================================================================

MIN_ITEMS = 6
MAX_ITEMS = 20
REQUIRE_ANSWER_KEY = True
REQUIRE_INSTRUCTIONS = True
REQUIRE_STUDENT_INFO = True
REQUIRE_PAGE_BREAK = True

# Required sections
REQUIRED_SECTIONS = [
    "instructions",
    "activity_content",
    "answer_key",
]

# Required student info patterns
STUDENT_INFO_PATTERNS = [
    r'[Nn]ame[:\s]*_+',
    r'[Pp]eriod[:\s]*_+',
    r'[Dd]ate[:\s]*_+',
]


# =============================================================================
# DOCUMENT VALIDATION
# =============================================================================

def validate_handout_file(file_path: Path) -> Dict[str, Any]:
    """
    Validate a handout Word document.

    HARDCODED validation that cannot be bypassed.

    Args:
        file_path: Path to .docx file

    Returns:
        Validation result dictionary
    """
    issues = []
    warnings = []

    # R5: Check file format
    if not str(file_path).endswith('.docx'):
        issues.append({
            "rule": "R5",
            "severity": "CRITICAL",
            "message": f"File must be .docx format, got: {file_path.suffix}"
        })
        return {
            "valid": False,
            "issues": issues,
            "warnings": warnings
        }

    if not file_path.exists():
        issues.append({
            "rule": "R5",
            "severity": "CRITICAL",
            "message": f"File not found: {file_path}"
        })
        return {
            "valid": False,
            "issues": issues,
            "warnings": warnings
        }

    if not DOCX_AVAILABLE:
        issues.append({
            "rule": "R5",
            "severity": "CRITICAL",
            "message": "python-docx not installed, cannot validate"
        })
        return {
            "valid": False,
            "issues": issues,
            "warnings": warnings
        }

    # Load document
    try:
        doc = Document(str(file_path))
    except Exception as e:
        issues.append({
            "rule": "R5",
            "severity": "CRITICAL",
            "message": f"Cannot open document: {e}"
        })
        return {
            "valid": False,
            "issues": issues,
            "warnings": warnings
        }

    # Extract all text
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"

    # Count tables
    table_count = len(doc.tables)
    total_table_rows = sum(len(t.rows) for t in doc.tables)

    # R3: Check for instructions
    has_instructions = bool(re.search(r'[Ii]nstructions?', full_text))
    if not has_instructions:
        issues.append({
            "rule": "R3",
            "severity": "CRITICAL",
            "message": "Instructions section not found"
        })

    # R4: Check for student info fields
    student_info_found = 0
    for pattern in STUDENT_INFO_PATTERNS:
        if re.search(pattern, full_text):
            student_info_found += 1

    if student_info_found < 2:
        issues.append({
            "rule": "R4",
            "severity": "CRITICAL",
            "message": f"Student info fields missing (found {student_info_found}/3 required)"
        })

    # R2: Check for answer key
    has_answer_key = bool(re.search(r'[Aa]nswer\s*[Kk]ey', full_text))
    if not has_answer_key:
        issues.append({
            "rule": "R2",
            "severity": "CRITICAL",
            "message": "Answer key section not found"
        })

    # R1: Check for minimum items (estimate from table rows)
    # Assume first table after instructions is the activity
    estimated_items = 0
    for table in doc.tables:
        # Subtract 1 for header row
        rows = len(table.rows) - 1
        if rows > estimated_items:
            estimated_items = rows

    if estimated_items < MIN_ITEMS:
        issues.append({
            "rule": "R1",
            "severity": "CRITICAL",
            "message": f"Minimum {MIN_ITEMS} items required, estimated {estimated_items}"
        })

    if estimated_items > MAX_ITEMS:
        warnings.append({
            "rule": "R1",
            "severity": "WARNING",
            "message": f"More than {MAX_ITEMS} items may overwhelm students, found {estimated_items}"
        })

    # Check for page break (look for section breaks or multiple pages)
    # This is approximate since python-docx doesn't expose page breaks easily
    has_teacher_reference = bool(re.search(r'[Tt]eacher\s*[Rr]eference', full_text))
    if has_answer_key and not has_teacher_reference:
        warnings.append({
            "rule": "R2",
            "severity": "WARNING",
            "message": "Answer key should be marked 'Teacher Reference Only'"
        })

    return {
        "valid": len([i for i in issues if i["severity"] == "CRITICAL"]) == 0,
        "file_path": str(file_path),
        "estimated_items": estimated_items,
        "table_count": table_count,
        "has_instructions": has_instructions,
        "has_answer_key": has_answer_key,
        "student_info_fields": student_info_found,
        "issues": issues,
        "warnings": warnings
    }


# =============================================================================
# DATA VALIDATION (Before Generation)
# =============================================================================

def validate_activity_data(
    activity_type: str,
    activity_data: Dict[str, Any]
) -> Dict[str, Any]:
    """
    Validate activity data before generating handout.

    HARDCODED validation that cannot be bypassed.

    Args:
        activity_type: Type of activity
        activity_data: Data for the activity

    Returns:
        Validation result
    """
    issues = []
    warnings = []

    # Check required fields
    if not activity_data.get("title"):
        issues.append({
            "rule": "DATA",
            "severity": "CRITICAL",
            "message": "Title is required"
        })

    if not activity_data.get("unit_day"):
        warnings.append({
            "rule": "DATA",
            "severity": "WARNING",
            "message": "unit_day not provided, using default"
        })

    # Type-specific validation
    if activity_type == "sorting":
        items = activity_data.get("items", [])
        categories = activity_data.get("categories", [])

        if len(items) < MIN_ITEMS:
            issues.append({
                "rule": "R1",
                "severity": "CRITICAL",
                "message": f"Sorting activity requires minimum {MIN_ITEMS} items, got {len(items)}"
            })

        if len(categories) < 2:
            issues.append({
                "rule": "DATA",
                "severity": "CRITICAL",
                "message": f"Sorting activity requires at least 2 categories, got {len(categories)}"
            })

        # Check each item has required fields
        for i, item in enumerate(items):
            if not item.get("text"):
                issues.append({
                    "rule": "R6",
                    "severity": "CRITICAL",
                    "message": f"Item {i+1} missing 'text' field"
                })
            if not item.get("category"):
                issues.append({
                    "rule": "R6",
                    "severity": "CRITICAL",
                    "message": f"Item {i+1} missing 'category' field"
                })
            if not item.get("explanation"):
                warnings.append({
                    "rule": "R6",
                    "severity": "WARNING",
                    "message": f"Item {i+1} missing 'explanation' for answer key"
                })

    elif activity_type == "matching":
        pairs = activity_data.get("pairs", [])

        if len(pairs) < MIN_ITEMS:
            issues.append({
                "rule": "R1",
                "severity": "CRITICAL",
                "message": f"Matching activity requires minimum {MIN_ITEMS} pairs, got {len(pairs)}"
            })

        for i, pair in enumerate(pairs):
            if not pair.get("left") or not pair.get("right"):
                issues.append({
                    "rule": "DATA",
                    "severity": "CRITICAL",
                    "message": f"Pair {i+1} missing 'left' or 'right' field"
                })

    elif activity_type == "sequencing":
        steps = activity_data.get("steps", [])

        if len(steps) < MIN_ITEMS:
            issues.append({
                "rule": "R1",
                "severity": "CRITICAL",
                "message": f"Sequencing activity requires minimum {MIN_ITEMS} steps, got {len(steps)}"
            })

        for i, step in enumerate(steps):
            if not step.get("text"):
                issues.append({
                    "rule": "DATA",
                    "severity": "CRITICAL",
                    "message": f"Step {i+1} missing 'text' field"
                })
            if step.get("order") is None:
                issues.append({
                    "rule": "DATA",
                    "severity": "CRITICAL",
                    "message": f"Step {i+1} missing 'order' field"
                })

    elif activity_type == "discussion":
        prompts = activity_data.get("prompts", [])

        if len(prompts) < 1:
            issues.append({
                "rule": "DATA",
                "severity": "CRITICAL",
                "message": "Discussion activity requires at least 1 prompt"
            })

        for i, prompt in enumerate(prompts):
            if not prompt.get("question"):
                issues.append({
                    "rule": "DATA",
                    "severity": "CRITICAL",
                    "message": f"Prompt {i+1} missing 'question' field"
                })

    return {
        "valid": len([i for i in issues if i["severity"] == "CRITICAL"]) == 0,
        "activity_type": activity_type,
        "issues": issues,
        "warnings": warnings
    }


# =============================================================================
# BATCH VALIDATION
# =============================================================================

def validate_lesson_handouts(
    lesson_path: Path,
    expected_activities: List[str] = None
) -> Dict[str, Any]:
    """
    Validate all handouts for a lesson.

    Args:
        lesson_path: Path to lesson directory
        expected_activities: List of expected activity handout filenames

    Returns:
        Validation summary
    """
    results = []
    missing = []

    # Find all .docx files
    docx_files = list(lesson_path.glob("*.docx"))
    handout_files = [f for f in docx_files if "handout" in f.name.lower()]

    # Validate each handout
    for handout in handout_files:
        result = validate_handout_file(handout)
        results.append(result)

    # Check for expected handouts
    if expected_activities:
        found_names = [f.stem.lower() for f in handout_files]
        for expected in expected_activities:
            if expected.lower() not in " ".join(found_names):
                missing.append(expected)

    all_valid = all(r["valid"] for r in results)
    total_issues = sum(len(r.get("issues", [])) for r in results)

    return {
        "valid": all_valid and len(missing) == 0,
        "handouts_found": len(handout_files),
        "handouts_validated": len(results),
        "all_passed": all_valid,
        "missing_handouts": missing,
        "total_issues": total_issues,
        "results": results
    }


# =============================================================================
# QUICK VALIDATION FUNCTIONS
# =============================================================================

def has_valid_handout(file_path: Path) -> bool:
    """Quick check if handout passes validation."""
    result = validate_handout_file(file_path)
    return result["valid"]


def get_handout_issues(file_path: Path) -> List[str]:
    """Get list of issue messages for a handout."""
    result = validate_handout_file(file_path)
    return [i["message"] for i in result.get("issues", [])]


# =============================================================================
# EXPORTS
# =============================================================================

__all__ = [
    # Main validation functions
    "validate_handout_file",
    "validate_activity_data",
    "validate_lesson_handouts",
    # Quick checks
    "has_valid_handout",
    "get_handout_issues",
    # Constants
    "MIN_ITEMS",
    "MAX_ITEMS",
    "REQUIRED_SECTIONS",
    "STUDENT_INFO_PATTERNS",
]

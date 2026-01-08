"""
File Parser - Parse .docx anchor documents for NCLEX pipeline.

This module provides utilities for parsing Word documents (.docx) containing
anchor point content for the NCLEX lecture generation pipeline.

Requirements:
- python-docx (pip install python-docx)

Usage:
    from skills.utilities.file_parser import FileParser

    parser = FileParser()
    content = parser.parse_docx("path/to/anchor_document.docx")
"""

import re
from pathlib import Path
from typing import Dict, List, Optional, Union
from dataclasses import dataclass, field


@dataclass
class ParsedDocument:
    """Container for parsed document content."""
    file_path: str
    file_name: str
    total_paragraphs: int
    total_words: int
    sections: List[Dict] = field(default_factory=list)
    raw_text: str = ""
    anchors: List[Dict] = field(default_factory=list)
    metadata: Dict = field(default_factory=dict)
    errors: List[str] = field(default_factory=list)


class FileParser:
    """Parse .docx anchor documents for the NCLEX pipeline."""

    # Common anchor patterns in documents
    ANCHOR_PATTERNS = [
        r'^(?:Anchor\s*)?(\d+)[\.:]\s*(.+)$',  # "Anchor 1: Topic" or "1. Topic"
        r'^\*\*\s*(.+?)\s*\*\*$',               # **Bold Topics**
        r'^#{1,3}\s*(.+)$',                     # Markdown headers
        r'^(?:Topic|Concept|Point)\s*\d*[\.:]\s*(.+)$',  # Topic/Concept markers
    ]

    def __init__(self, config_path: Optional[str] = None):
        """
        Initialize the FileParser.

        Args:
            config_path: Optional path to configuration file
        """
        self.config_path = config_path
        self._docx_available = self._check_docx_available()

    def _check_docx_available(self) -> bool:
        """Check if python-docx is available."""
        try:
            import docx
            return True
        except ImportError:
            return False

    def parse_docx(self, file_path: str) -> ParsedDocument:
        """
        Parse a .docx file and extract structured content.

        Args:
            file_path: Path to the .docx file

        Returns:
            ParsedDocument with extracted content
        """
        path = Path(file_path)
        result = ParsedDocument(
            file_path=str(path),
            file_name=path.name,
            total_paragraphs=0,
            total_words=0
        )

        # Validate file exists
        if not path.exists():
            result.errors.append(f"File not found: {file_path}")
            return result

        # Validate file extension
        if path.suffix.lower() != '.docx':
            result.errors.append(f"Invalid file format: {path.suffix}. Expected .docx")
            return result

        # Check if python-docx is available
        if not self._docx_available:
            result.errors.append("python-docx not installed. Run: pip install python-docx")
            return result

        try:
            from docx import Document
            doc = Document(file_path)

            # Extract all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append({
                        'text': text,
                        'style': para.style.name if para.style else 'Normal',
                        'is_heading': 'Heading' in (para.style.name if para.style else '')
                    })

            result.total_paragraphs = len(paragraphs)
            result.raw_text = '\n'.join([p['text'] for p in paragraphs])
            result.total_words = len(result.raw_text.split())

            # Extract anchors from paragraphs
            result.anchors = self._extract_anchors(paragraphs)

            # Extract document sections (based on headings)
            result.sections = self._extract_sections(paragraphs)

            # Extract metadata
            result.metadata = self._extract_metadata(doc, path)

        except Exception as e:
            result.errors.append(f"Parse error: {str(e)}")

        return result

    def parse_text(self, text: str, source_name: str = "text_input") -> ParsedDocument:
        """
        Parse raw text content (for testing or non-docx sources).

        Args:
            text: Raw text content
            source_name: Name to use for the source

        Returns:
            ParsedDocument with extracted content
        """
        result = ParsedDocument(
            file_path=source_name,
            file_name=source_name,
            total_paragraphs=0,
            total_words=0
        )

        lines = text.strip().split('\n')
        paragraphs = []

        for line in lines:
            text_line = line.strip()
            if text_line:
                paragraphs.append({
                    'text': text_line,
                    'style': 'Normal',
                    'is_heading': text_line.startswith('#') or text_line.isupper()
                })

        result.total_paragraphs = len(paragraphs)
        result.raw_text = '\n'.join([p['text'] for p in paragraphs])
        result.total_words = len(result.raw_text.split())
        result.anchors = self._extract_anchors(paragraphs)
        result.sections = self._extract_sections(paragraphs)

        return result

    def _extract_anchors(self, paragraphs: List[Dict]) -> List[Dict]:
        """
        Extract anchor points from paragraphs.

        Args:
            paragraphs: List of paragraph dictionaries

        Returns:
            List of anchor dictionaries
        """
        anchors = []
        anchor_id = 0

        for para in paragraphs:
            text = para['text']

            # Check for explicit anchor patterns
            for pattern in self.ANCHOR_PATTERNS:
                match = re.match(pattern, text, re.IGNORECASE)
                if match:
                    anchor_id += 1
                    groups = match.groups()

                    # Determine anchor title from matched groups
                    if len(groups) == 2:
                        # Pattern with number and title
                        title = groups[1].strip()
                    else:
                        # Pattern with just title
                        title = groups[0].strip()

                    anchors.append({
                        'id': f"anchor_{anchor_id}",
                        'number': anchor_id,
                        'title': title,
                        'full_text': text,
                        'style': para.get('style', 'Normal'),
                        'is_heading': para.get('is_heading', False)
                    })
                    break
            else:
                # Check if it's a heading that could be an anchor
                if para.get('is_heading') or para.get('style', '').startswith('Heading'):
                    anchor_id += 1
                    anchors.append({
                        'id': f"anchor_{anchor_id}",
                        'number': anchor_id,
                        'title': text,
                        'full_text': text,
                        'style': para.get('style', 'Normal'),
                        'is_heading': True
                    })

        return anchors

    def _extract_sections(self, paragraphs: List[Dict]) -> List[Dict]:
        """
        Extract document sections based on headings.

        Args:
            paragraphs: List of paragraph dictionaries

        Returns:
            List of section dictionaries
        """
        sections = []
        current_section = None

        for para in paragraphs:
            if para.get('is_heading') or para.get('style', '').startswith('Heading'):
                # Start new section
                if current_section:
                    sections.append(current_section)

                current_section = {
                    'title': para['text'],
                    'style': para.get('style', 'Normal'),
                    'content': []
                }
            elif current_section:
                current_section['content'].append(para['text'])

        # Add last section
        if current_section:
            sections.append(current_section)

        return sections

    def _extract_metadata(self, doc, path: Path) -> Dict:
        """
        Extract document metadata.

        Args:
            doc: python-docx Document object
            path: Path to the document

        Returns:
            Dictionary of metadata
        """
        metadata = {
            'file_size_bytes': path.stat().st_size if path.exists() else 0,
            'file_modified': path.stat().st_mtime if path.exists() else None,
        }

        try:
            core_props = doc.core_properties
            metadata.update({
                'author': core_props.author or '',
                'title': core_props.title or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
            })
        except Exception:
            pass

        return metadata

    def validate_parsed_document(self, parsed: ParsedDocument) -> Dict:
        """
        Validate a parsed document meets minimum requirements.

        Args:
            parsed: ParsedDocument to validate

        Returns:
            Dictionary with validation results
        """
        validation = {
            'is_valid': True,
            'warnings': [],
            'errors': parsed.errors.copy()
        }

        # Check for parse errors
        if parsed.errors:
            validation['is_valid'] = False

        # Check minimum anchor count
        if len(parsed.anchors) < 5:
            validation['warnings'].append(
                f"Low anchor count: {len(parsed.anchors)} (recommend >= 20)"
            )

        if len(parsed.anchors) == 0:
            validation['is_valid'] = False
            validation['errors'].append("No anchors found in document")

        # Check word count
        if parsed.total_words < 100:
            validation['warnings'].append(
                f"Low word count: {parsed.total_words}"
            )

        return validation


def parse_file(file_path: str) -> ParsedDocument:
    """
    Convenience function to parse a file.

    Args:
        file_path: Path to file to parse

    Returns:
        ParsedDocument with extracted content
    """
    parser = FileParser()
    return parser.parse_docx(file_path)


if __name__ == "__main__":
    import sys

    print("File Parser - NCLEX Pipeline Utility")
    print("=" * 50)

    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        print(f"Parsing: {file_path}")
        print()

        parser = FileParser()
        result = parser.parse_docx(file_path)

        if result.errors:
            print("ERRORS:")
            for error in result.errors:
                print(f"  - {error}")
        else:
            print(f"File: {result.file_name}")
            print(f"Paragraphs: {result.total_paragraphs}")
            print(f"Words: {result.total_words}")
            print(f"Anchors found: {len(result.anchors)}")
            print(f"Sections found: {len(result.sections)}")
            print()

            if result.anchors:
                print("First 5 anchors:")
                for anchor in result.anchors[:5]:
                    print(f"  {anchor['number']}. {anchor['title'][:50]}")

            # Validate
            validation = parser.validate_parsed_document(result)
            print()
            print(f"Validation: {'PASS' if validation['is_valid'] else 'FAIL'}")
            if validation['warnings']:
                print("Warnings:")
                for warn in validation['warnings']:
                    print(f"  - {warn}")
    else:
        # Demo with sample text
        sample_text = """
        # Introduction to NCLEX Preparation

        Anchor 1: Assessment Fundamentals
        Understanding patient assessment is critical for nursing practice.

        Anchor 2: Vital Signs Monitoring
        Accurate vital sign measurement and interpretation.

        Anchor 3: Documentation Standards
        Proper charting and documentation requirements.
        """

        print("Demo mode - parsing sample text")
        print()

        parser = FileParser()
        result = parser.parse_text(sample_text, "demo_input")

        print(f"Paragraphs: {result.total_paragraphs}")
        print(f"Words: {result.total_words}")
        print(f"Anchors found: {len(result.anchors)}")
        print()
        print("Anchors:")
        for anchor in result.anchors:
            print(f"  {anchor['number']}. {anchor['title']}")
